import streamlit as st
import openai
import json
from datetime import datetime
import pandas as pd
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import spacy
from collections import Counter
import re
import sqlite3
import hashlib
import secrets

# Page configuration
st.set_page_config(
    page_title="AI Marketing Content Generator",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .sub-header {
        font-size: 1.5rem;
        color: #ff7f0e;
        margin-top: 2rem;
    }
    .content-box {
        background-color: #f0f2f6;
        padding: 20px;
        border-radius: 10px;
        border-left: 5px solid #1f77b4;
        margin: 10px 0;
    }
    .keyword-badge {
        display: inline-block;
        padding: 5px 10px;
        margin: 5px;
        background-color: #1f77b4;
        color: white;
        border-radius: 15px;
        font-size: 0.9rem;
    }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 20px;
        border-radius: 10px;
        color: white;
        text-align: center;
    }
    .success-message {
        background-color: #d4edda;
        color: #155724;
        padding: 15px;
        border-radius: 5px;
        border: 1px solid #c3e6cb;
    }
    .stButton>button {
        width: 100%;
        background-color: #1f77b4;
        color: white;
        border-radius: 5px;
        padding: 10px;
        font-weight: bold;
    }
</style>
""", unsafe_allow_html=True)

# Database initialization
def init_db():
    conn = sqlite3.connect('marketing_content.db', check_same_thread=False)
    c = conn.cursor()
    
    # Users table
    c.execute('''CREATE TABLE IF NOT EXISTS users
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  username TEXT UNIQUE NOT NULL,
                  email TEXT UNIQUE NOT NULL,
                  password_hash TEXT NOT NULL,
                  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
    
    # Generated content table
    c.execute('''CREATE TABLE IF NOT EXISTS generated_content
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  user_id INTEGER,
                  platform TEXT NOT NULL,
                  product_name TEXT NOT NULL,
                  product_description TEXT,
                  target_audience TEXT,
                  brand_tone TEXT,
                  keywords TEXT,
                  headline TEXT,
                  body_content TEXT,
                  cta TEXT,
                  hashtags TEXT,
                  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                  FOREIGN KEY (user_id) REFERENCES users (id))''')
    
    # Sessions table
    c.execute('''CREATE TABLE IF NOT EXISTS sessions
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  user_id INTEGER,
                  session_token TEXT UNIQUE,
                  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                  expires_at TIMESTAMP,
                  FOREIGN KEY (user_id) REFERENCES users (id))''')
    
    conn.commit()
    conn.close()

# Initialize database
init_db()

# Database helper functions
def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def create_user(username, email, password):
    try:
        conn = sqlite3.connect('marketing_content.db', check_same_thread=False)
        c = conn.cursor()
        password_hash = hash_password(password)
        c.execute("INSERT INTO users (username, email, password_hash) VALUES (?, ?, ?)",
                  (username, email, password_hash))
        conn.commit()
        conn.close()
        return True
    except sqlite3.IntegrityError:
        return False

def verify_user(username, password):
    conn = sqlite3.connect('marketing_content.db', check_same_thread=False)
    c = conn.cursor()
    password_hash = hash_password(password)
    c.execute("SELECT id, username FROM users WHERE username=? AND password_hash=?",
              (username, password_hash))
    user = c.fetchone()
    conn.close()
    return user

def save_content_to_db(user_id, platform, product_name, description, audience, tone, 
                       keywords, headline, body, cta, hashtags):
    conn = sqlite3.connect('marketing_content.db', check_same_thread=False)
    c = conn.cursor()
    c.execute("""INSERT INTO generated_content 
                 (user_id, platform, product_name, product_description, target_audience, 
                  brand_tone, keywords, headline, body_content, cta, hashtags)
                 VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
              (user_id, platform, product_name, description, audience, tone,
               json.dumps(keywords), headline, body, cta, json.dumps(hashtags)))
    conn.commit()
    conn.close()

def get_user_content_history(user_id, limit=50):
    conn = sqlite3.connect('marketing_content.db', check_same_thread=False)
    c = conn.cursor()
    c.execute("""SELECT id, platform, product_name, headline, body_content, 
                 cta, hashtags, created_at FROM generated_content 
                 WHERE user_id=? ORDER BY created_at DESC LIMIT ?""", (user_id, limit))
    content = c.fetchall()
    conn.close()
    return content

def delete_content(content_id, user_id):
    conn = sqlite3.connect('marketing_content.db', check_same_thread=False)
    c = conn.cursor()
    c.execute("DELETE FROM generated_content WHERE id=? AND user_id=?", (content_id, user_id))
    conn.commit()
    conn.close()

# Load spaCy model
@st.cache_resource
def load_spacy_model():
    try:
        nlp = spacy.load("en_core_web_sm")
    except:
        import os
        os.system("python -m spacy download en_core_web_sm")
        nlp = spacy.load("en_core_web_sm")
    return nlp

nlp = load_spacy_model()

# NLP Keyword Extraction
def extract_keywords(text, top_n=10):
    doc = nlp(text.lower())
    
    # Extract nouns and adjectives
    keywords = []
    for token in doc:
        if token.pos_ in ['NOUN', 'ADJ', 'PROPN'] and not token.is_stop and len(token.text) > 2:
            keywords.append(token.lemma_)
    
    # Count frequency
    keyword_freq = Counter(keywords)
    
    # Get top keywords
    top_keywords = [word for word, freq in keyword_freq.most_common(top_n)]
    
    return top_keywords

# Platform-specific prompts
PLATFORM_PROMPTS = {
    "Google Ads": {
        "system": "You are an expert Google Ads copywriter specializing in high-converting ad copy.",
        "instructions": """Generate Google Ads copy with:
        - Headline (max 30 characters)
        - Description (max 90 characters)
        - Strong CTA
        - Keyword optimization
        
        Focus on benefits, urgency, and relevance score optimization."""
    },
    "Facebook Ads": {
        "system": "You are an expert Facebook advertising specialist.",
        "instructions": """Generate Facebook Ad copy with:
        - Attention-grabbing hook
        - Engaging body text (125-150 words)
        - Emotional appeal
        - Clear CTA
        - 3-5 relevant hashtags
        
        Use conversational tone and address pain points directly."""
    },
    "Instagram": {
        "system": "You are an Instagram marketing expert specializing in visual storytelling.",
        "instructions": """Generate Instagram post copy with:
        - Captivating first line
        - Story-driven content (150-200 words)
        - Emoji integration
        - Strong CTA
        - 10-15 trending hashtags
        
        Focus on visual language and community engagement."""
    },
    "SEO Meta Description": {
        "system": "You are an SEO specialist focused on search engine optimization.",
        "instructions": """Generate SEO-optimized meta description with:
        - Compelling description (150-160 characters)
        - Primary keyword integration
        - Benefit-focused language
        - CTA or value proposition
        
        Optimize for click-through rate and search relevance."""
    },
    "Landing Page": {
        "system": "You are a conversion-focused landing page copywriter.",
        "instructions": """Generate landing page content with:
        - Powerful headline
        - Subheadline
        - 3-4 benefit bullet points
        - Social proof statement
        - Primary and secondary CTA
        
        Focus on conversion optimization and value proposition."""
    }
}

# AI Content Generation
def generate_content(api_key, platform, product_name, description, audience, tone, keywords):
    openai.api_key = api_key
    
    platform_config = PLATFORM_PROMPTS[platform]
    
    prompt = f"""
{platform_config['instructions']}

Product Information:
- Name: {product_name}
- Description: {description}
- Target Audience: {audience}
- Brand Tone: {tone}
- Keywords: {', '.join(keywords)}

Return the response in the following JSON format:
{{
    "headline": "compelling headline",
    "body": "main content body",
    "cta": "call to action",
    "hashtags": ["hashtag1", "hashtag2", "hashtag3"]
}}
"""
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": platform_config['system']},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=800
        )
        
        content = response.choices[0].message.content
        
        # Parse JSON response
        try:
            content_json = json.loads(content)
        except:
            # Fallback if not JSON
            content_json = {
                "headline": "Check out our amazing product!",
                "body": content,
                "cta": "Shop Now",
                "hashtags": keywords[:5]
            }
        
        return content_json
    
    except Exception as e:
        st.error(f"Error generating content: {str(e)}")
        return None

# Export to PDF
def generate_pdf(content_data):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1f77b4'),
        spaceAfter=30,
        alignment=1
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#ff7f0e'),
        spaceAfter=12,
        spaceBefore=12
    )
    
    # Title
    story.append(Paragraph("AI Marketing Content Report", title_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Metadata
    data = [
        ['Platform:', content_data['platform']],
        ['Product:', content_data['product_name']],
        ['Generated:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
        ['Tone:', content_data['tone']],
        ['Audience:', content_data['audience']]
    ]
    
    table = Table(data, colWidths=[2*inch, 4*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f2f6')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey)
    ]))
    
    story.append(table)
    story.append(Spacer(1, 0.3*inch))
    
    # Content sections
    story.append(Paragraph("Headline", heading_style))
    story.append(Paragraph(content_data['headline'], styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    story.append(Paragraph("Body Content", heading_style))
    story.append(Paragraph(content_data['body'], styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    story.append(Paragraph("Call to Action", heading_style))
    story.append(Paragraph(content_data['cta'], styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    if content_data['hashtags']:
        story.append(Paragraph("Hashtags", heading_style))
        hashtags_text = ' '.join([f"#{tag}" for tag in content_data['hashtags']])
        story.append(Paragraph(hashtags_text, styles['Normal']))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

# Export to DOCX
def generate_docx(content_data):
    doc = Document()
    
    # Title
    title = doc.add_heading('AI Marketing Content Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Platform: {content_data['platform']}")
    doc.add_paragraph(f"Product: {content_data['product_name']}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Tone: {content_data['tone']}")
    doc.add_paragraph(f"Audience: {content_data['audience']}")
    
    doc.add_paragraph()
    
    # Headline
    doc.add_heading('Headline', level=1)
    p = doc.add_paragraph(content_data['headline'])
    p.runs[0].font.size = Pt(14)
    
    # Body
    doc.add_heading('Body Content', level=1)
    doc.add_paragraph(content_data['body'])
    
    # CTA
    doc.add_heading('Call to Action', level=1)
    p = doc.add_paragraph(content_data['cta'])
    p.runs[0].bold = True
    
    # Hashtags
    if content_data['hashtags']:
        doc.add_heading('Hashtags', level=1)
        hashtags_text = ' '.join([f"#{tag}" for tag in content_data['hashtags']])
        doc.add_paragraph(hashtags_text)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Authentication
def login_page():
    st.markdown("<h1 class='main-header'>🚀 AI Marketing Content Generator</h1>", unsafe_allow_html=True)
    
    tab1, tab2 = st.tabs(["Login", "Sign Up"])
    
    with tab1:
        st.subheader("Login to Your Account")
        username = st.text_input("Username", key="login_username")
        password = st.text_input("Password", type="password", key="login_password")
        
        if st.button("Login", key="login_btn"):
            if username and password:
                user = verify_user(username, password)
                if user:
                    st.session_state['logged_in'] = True
                    st.session_state['user_id'] = user[0]
                    st.session_state['username'] = user[1]
                    st.success("Login successful!")
                    st.rerun()
                else:
                    st.error("Invalid username or password")
            else:
                st.warning("Please enter both username and password")
    
    with tab2:
        st.subheader("Create New Account")
        new_username = st.text_input("Username", key="signup_username")
        new_email = st.text_input("Email", key="signup_email")
        new_password = st.text_input("Password", type="password", key="signup_password")
        confirm_password = st.text_input("Confirm Password", type="password", key="confirm_password")
        
        if st.button("Sign Up", key="signup_btn"):
            if new_username and new_email and new_password and confirm_password:
                if new_password != confirm_password:
                    st.error("Passwords do not match")
                elif len(new_password) < 6:
                    st.error("Password must be at least 6 characters")
                else:
                    if create_user(new_username, new_email, new_password):
                        st.success("Account created successfully! Please login.")
                    else:
                        st.error("Username or email already exists")
            else:
                st.warning("Please fill all fields")

# Main application
def main_app():
    # Sidebar
    with st.sidebar:
        st.image("https://via.placeholder.com/150x50?text=AI+Marketing", use_column_width=True)
        st.markdown(f"### Welcome, {st.session_state['username']}!")
        
        if st.button("Logout"):
            st.session_state['logged_in'] = False
            st.rerun()
        
        st.markdown("---")
        
        page = st.radio("Navigation", ["Generate Content", "Content History", "Analytics", "Keyword Extractor"])
    
    # Main content
    if page == "Generate Content":
        generate_content_page()
    elif page == "Content History":
        content_history_page()
    elif page == "Analytics":
        analytics_page()
    elif page == "Keyword Extractor":
        keyword_extractor_page()

def generate_content_page():
    st.markdown("<h1 class='main-header'>Generate Marketing Content</h1>", unsafe_allow_html=True)
    
    # API Key input
    api_key = st.text_input("OpenAI API Key", type="password", help="Enter your OpenAI API key")
    
    if not api_key:
        st.warning("⚠️ Please enter your OpenAI API key to generate content")
        st.info("Get your API key from: https://platform.openai.com/api-keys")
        return
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("Product Information")
        product_name = st.text_input("Product Name", placeholder="e.g., EcoFriendly Water Bottle")
        product_description = st.text_area("Product Description", 
                                          placeholder="Describe your product features, benefits, and unique selling points...",
                                          height=150)
        
        target_audience = st.text_input("Target Audience", 
                                       placeholder="e.g., Environmentally conscious millennials")
    
    with col2:
        st.subheader("Content Settings")
        platform = st.selectbox("Platform", list(PLATFORM_PROMPTS.keys()))
        brand_tone = st.selectbox("Brand Tone", 
                                 ["Professional", "Casual", "Witty", "Urgent", "Inspirational", "Friendly"])
    
    # Auto-extract keywords
    if product_description:
        with st.expander("🔍 Auto-Extracted Keywords", expanded=False):
            keywords = extract_keywords(product_description)
            st.write("Top keywords from your description:")
            keyword_html = "".join([f"<span class='keyword-badge'>{kw}</span>" for kw in keywords])
            st.markdown(keyword_html, unsafe_allow_html=True)
    else:
        keywords = []
    
    st.markdown("---")
    
    if st.button("🚀 Generate Content", type="primary", use_container_width=True):
        if not product_name or not product_description:
            st.error("Please fill in product name and description")
            return
        
        with st.spinner("🤖 AI is crafting your perfect content..."):
            content = generate_content(api_key, platform, product_name, product_description, 
                                     target_audience, brand_tone, keywords)
            
            if content:
                st.session_state['generated_content'] = content
                st.session_state['content_metadata'] = {
                    'platform': platform,
                    'product_name': product_name,
                    'product_description': product_description,
                    'target_audience': target_audience,
                    'tone': brand_tone,
                    'keywords': keywords
                }
                
                # Save to database
                save_content_to_db(
                    st.session_state['user_id'],
                    platform,
                    product_name,
                    product_description,
                    target_audience,
                    brand_tone,
                    keywords,
                    content['headline'],
                    content['body'],
                    content['cta'],
                    content.get('hashtags', [])
                )
                
                st.success("✅ Content generated successfully!")
    
    # Display generated content
    if 'generated_content' in st.session_state:
        st.markdown("---")
        st.markdown("<h2 class='sub-header'>Generated Content</h2>", unsafe_allow_html=True)
        
        content = st.session_state['generated_content']
        metadata = st.session_state['content_metadata']
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown("<div class='content-box'><strong>📱 Platform</strong><br>" + 
                       metadata['platform'] + "</div>", unsafe_allow_html=True)
        with col2:
            st.markdown("<div class='content-box'><strong>🎯 Audience</strong><br>" + 
                       metadata['target_audience'] + "</div>", unsafe_allow_html=True)
        with col3:
            st.markdown("<div class='content-box'><strong>🎨 Tone</strong><br>" + 
                       metadata['tone'] + "</div>", unsafe_allow_html=True)
        
        st.markdown("### 📝 Headline")
        st.markdown(f"<div class='content-box'><h3>{content['headline']}</h3></div>", 
                   unsafe_allow_html=True)
        
        st.markdown("### 📄 Body Content")
        st.markdown(f"<div class='content-box'>{content['body']}</div>", unsafe_allow_html=True)
        
        st.markdown("### 🎯 Call to Action")
        st.markdown(f"<div class='content-box'><strong>{content['cta']}</strong></div>", 
                   unsafe_allow_html=True)
        
        if content.get('hashtags'):
            st.markdown("### #️⃣ Hashtags")
            hashtag_html = "".join([f"<span class='keyword-badge'>#{tag}</span>" 
                                   for tag in content['hashtags']])
            st.markdown(hashtag_html, unsafe_allow_html=True)
        
        # Export options
        st.markdown("---")
        col1, col2 = st.columns(2)
        
        with col1:
            pdf_buffer = generate_pdf({
                'platform': metadata['platform'],
                'product_name': metadata['product_name'],
                'tone': metadata['tone'],
                'audience': metadata['target_audience'],
                'headline': content['headline'],
                'body': content['body'],
                'cta': content['cta'],
                'hashtags': content.get('hashtags', [])
            })
            st.download_button(
                label="📥 Download as PDF",
                data=pdf_buffer,
                file_name=f"{metadata['product_name']}_content.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        
        with col2:
            docx_buffer = generate_docx({
                'platform': metadata['platform'],
                'product_name': metadata['product_name'],
                'tone': metadata['tone'],
                'audience': metadata['target_audience'],
                'headline': content['headline'],
                'body': content['body'],
                'cta': content['cta'],
                'hashtags': content.get('hashtags', [])
            })
            st.download_button(
                label="📥 Download as DOCX",
                data=docx_buffer,
                file_name=f"{metadata['product_name']}_content.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

def content_history_page():
    st.markdown("<h1 class='main-header'>Content History</h1>", unsafe_allow_html=True)
    
    history = get_user_content_history(st.session_state['user_id'])
    
    if not history:
        st.info("No content generated yet. Start creating amazing content!")
        return
    
    st.markdown(f"**Total Content Generated: {len(history)}**")
    
    for item in history:
        with st.expander(f"{item[1]} - {item[2]} | {item[7]}"):
            col1, col2 = st.columns([3, 1])
            
            with col1:
                st.markdown(f"**Headline:** {item[3]}")
                st.markdown(f"**Body:** {item[4]}")
                st.markdown(f"**CTA:** {item[5]}")
                if item[6]:
                    hashtags = json.loads(item[6])
                    st.markdown(f"**Hashtags:** {' '.join([f'#{h}' for h in hashtags])}")
            
            with col2:
                if st.button("🗑️ Delete", key=f"del_{item[0]}"):
                    delete_content(item[0], st.session_state['user_id'])
                    st.success("Deleted!")
                    st.rerun()

def analytics_page():
    st.markdown("<h1 class='main-header'>Analytics Dashboard</h1>", unsafe_allow_html=True)
    
    history = get_user_content_history(st.session_state['user_id'], limit=1000)
    
    if not history:
        st.info("Generate some content to see analytics!")
        return
    
    # Convert to DataFrame
    df = pd.DataFrame(history, columns=['id', 'platform', 'product', 'headline', 
                                        'body', 'cta', 'hashtags', 'created_at'])
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("<div class='metric-card'>", unsafe_allow_html=True)
        st.metric("Total Content", len(df))
        st.markdown("</div>", unsafe_allow_html=True)
    
    with col2:
        st.markdown("<div class='metric-card'>", unsafe_allow_html=True)
        st.metric("Platforms Used", df['platform'].nunique())
        st.markdown("</div>", unsafe_allow_html=True)
    
    with col3:
        st.markdown("<div class='metric-card'>", unsafe_allow_html=True)
        st.metric("Products", df['product'].nunique())
        st.markdown("</div>", unsafe_allow_html=True)
    
    st.markdown("---")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Content by Platform")
        platform_counts = df['platform'].value_counts()
        st.bar_chart(platform_counts)
    
    with col2:
        st.subheader("Platform Distribution")
        st.dataframe(platform_counts, use_container_width=True)

def keyword_extractor_page():
    st.markdown("<h1 class='main-header'>Keyword Extractor</h1>", unsafe_allow_html=True)
    
    st.markdown("""
    Extract relevant keywords from your product descriptions using advanced NLP technology.
    """)
    
    text_input = st.text_area("Enter text to extract keywords", height=200,
                             placeholder="Paste your product description or marketing copy here...")
    
    top_n = st.slider("Number of keywords to extract", 5, 20, 10)
    
    if st.button("Extract Keywords", type="primary"):
        if text_input:
            with st.spinner("Extracting keywords..."):
                keywords = extract_keywords(text_input, top_n)
                
                st.success(f"Extracted {len(keywords)} keywords!")
                
                # Display as badges
                keyword_html = "".join([f"<span class='keyword-badge'>{kw}</span>" 
                                       for kw in keywords])
                st.markdown(keyword_html, unsafe_allow_html=True)
                
                # Display as list
                st.markdown("---")
                st.subheader("Keyword List (copy-paste ready)")
                st.code(", ".join(keywords))
        else:
            st.warning("Please enter some text to extract keywords")

# Main execution
def main():
    if 'logged_in' not in st.session_state:
        st.session_state['logged_in'] = False
    
    if not st.session_state['logged_in']:
        login_page()
    else:
        main_app()

if __name__ == "__main__":
    main()